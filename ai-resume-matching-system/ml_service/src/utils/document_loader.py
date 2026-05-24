# =========================================================
# document_loader.py
# =========================================================

# =========================================================
# IMPORTS
# =========================================================

# PyMuPDF library
# used for PDF text extraction

import fitz


# python-docx library
# used for DOCX text extraction

from docx import Document


# os library
# used for file extension checking

import os



# =========================================================
# PDF TEXT EXTRACTION
# =========================================================

def extract_text_from_pdf(pdf_path):

    """
    Extract text from PDF resume.

    Parameters:
    pdf_path (str):
        path to uploaded pdf file

    Returns:
    str:
        extracted full text
    """

    # -----------------------------------------------------
    # OPEN PDF DOCUMENT
    # -----------------------------------------------------

    # fitz.open(...)
    # loads pdf into memory

    doc = fitz.open(pdf_path)


    # -----------------------------------------------------
    # EMPTY STRING
    # -----------------------------------------------------

    # We accumulate all page text here

    text = ""


    # -----------------------------------------------------
    # LOOP THROUGH PAGES
    # -----------------------------------------------------

    # PDFs may contain multiple pages

    for page in doc:


        # -------------------------------------------------
        # EXTRACT PAGE TEXT
        # -------------------------------------------------

        # page.get_text()
        # extracts readable text from current page

        text += page.get_text()


    # -----------------------------------------------------
    # RETURN FINAL TEXT
    # -----------------------------------------------------

    return text



# =========================================================
# DOCX TEXT EXTRACTION
# =========================================================

def extract_text_from_docx(docx_path):

    """
    Extract text from DOCX resume.

    Parameters:
    docx_path (str):
        path to uploaded docx file

    Returns:
    str:
        extracted full text
    """

    # -----------------------------------------------------
    # LOAD DOCX DOCUMENT
    # -----------------------------------------------------

    # Document(...)
    # loads Word file

    doc = Document(docx_path)


    # -----------------------------------------------------
    # EXTRACT ALL PARAGRAPHS
    # -----------------------------------------------------

    # doc.paragraphs
    # gives all paragraph objects

    # para.text
    # extracts paragraph text

    # "\n".join(...)
    # combines paragraphs into one string

    text = "\n".join(
        [para.text for para in doc.paragraphs]
    )


    # -----------------------------------------------------
    # RETURN FINAL TEXT
    # -----------------------------------------------------

    return text



# =========================================================
# GENERAL FILE LOADER
# =========================================================

def extract_resume_text(file_path):

    """
    Automatically detect resume type
    and extract text.

    Supported:
    - PDF
    - DOCX

    Parameters:
    file_path (str)

    Returns:
    str:
        extracted resume text
    """

    # -----------------------------------------------------
    # GET FILE EXTENSION
    # -----------------------------------------------------

    # os.path.splitext(...)
    # separates:
    # filename + extension

    # Example:
    # resume.pdf
    # becomes:
    # ('resume', '.pdf')

    _, extension = os.path.splitext(file_path)


    # -----------------------------------------------------
    # LOWERCASE EXTENSION
    # -----------------------------------------------------

    # Makes comparison safer

    extension = extension.lower()


    # -----------------------------------------------------
    # PDF HANDLING
    # -----------------------------------------------------

    if extension == ".pdf":

        return extract_text_from_pdf(file_path)


    # -----------------------------------------------------
    # DOCX HANDLING
    # -----------------------------------------------------

    elif extension == ".docx":

        return extract_text_from_docx(file_path)


    # -----------------------------------------------------
    # UNSUPPORTED FILE
    # -----------------------------------------------------

    else:

        raise ValueError(
            "Unsupported file format. Use PDF or DOCX."
        )



# =========================================================
# OPTIONAL TESTING
# =========================================================

if __name__ == "__main__":

    # Example test file

    sample_file = "../../uploads/sample_resume.pdf"


    # Extract text

    extracted_text = extract_resume_text(
        sample_file
    )


    # Print extracted text

    print(extracted_text)
